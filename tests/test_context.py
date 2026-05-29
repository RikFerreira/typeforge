import docx
import pytest
from docx import Document

from typeforge.jinja_context_builder.context import Context


# --- Construction -----------------------------------------------------------

def test_single_dict_and_str_label_are_wrapped():
    ctx = Context({"a": 1}, "formulario")

    assert ctx._context == {"formulario": {"a": 1}}


def test_lists_are_zipped_by_position():
    data = [{"a": 1}, {"b": 2}]
    labels = ["first", "second"]

    ctx = Context(data, labels)

    assert ctx._context == {"first": {"a": 1}, "second": {"b": 2}}


def test_mismatched_lengths_raise():
    with pytest.raises(Exception, match="length"):
        Context([{"a": 1}, {"b": 2}], ["only_one"])


def test_non_string_label_raises():
    with pytest.raises(Exception, match="strings"):
        Context([{"a": 1}], [123])


def test_repr_contains_context():
    ctx = Context({"a": 1}, "label")

    assert "Context(dict_context=" in repr(ctx)


# --- Filter registration ----------------------------------------------------

def test_filters_module_is_auto_registered():
    ctx = Context({"a": 1}, "label")

    # Every public function in filters.py should be available as a jinja filter.
    assert "pluck_first" in ctx._jinja_env.filters
    assert "pluck_where" in ctx._jinja_env.filters


def test_add_filter_makes_it_available():
    ctx = Context({"a": 1}, "label")

    ctx.add_filter("shout", lambda s: str(s).upper())

    assert ctx._jinja_env.filters["shout"]("hi") == "HI"


def test_add_duplicate_filter_raises():
    ctx = Context({"a": 1}, "label")

    with pytest.raises(Exception, match="already exists"):
        ctx.add_filter("pluck_first", lambda x: x)


def test_remove_filter_drops_it():
    ctx = Context({"a": 1}, "label")

    ctx.remove_filter("pluck_first")

    assert "pluck_first" not in ctx._jinja_env.filters


def test_remove_missing_filter_warns_and_returns_none():
    ctx = Context({"a": 1}, "label")

    with pytest.warns(UserWarning, match="No filter"):
        result = ctx.remove_filter("does_not_exist")

    assert result is None


# --- Rendering --------------------------------------------------------------

def _make_template(tmp_path, text):
    """Build a minimal .docx template containing a single jinja paragraph."""
    template_path = tmp_path / "template.docx"
    document = Document()
    document.add_paragraph(text)
    document.save(template_path)
    return template_path


def _read_paragraphs(path):
    return [p.text for p in Document(path).paragraphs]


def test_render_docx_substitutes_context(tmp_path):
    template_path = _make_template(tmp_path, "Hello {{ greeting.name }}")
    output_path = tmp_path / "out.docx"

    ctx = Context({"name": "World"}, "greeting")
    ctx.render_docx(str(template_path), str(output_path))

    assert output_path.exists()
    assert "Hello World" in _read_paragraphs(output_path)


def test_render_docx_applies_registered_filter(tmp_path):
    template_path = _make_template(
        tmp_path, "{{ data | pluck_first('city') }}"
    )
    output_path = tmp_path / "out.docx"

    ctx = Context(
        [[{"city": "Belem"}, {"city": "Manaus"}]],
        ["data"],
    )
    ctx.render_docx(str(template_path), str(output_path))

    assert "Belem" in _read_paragraphs(output_path)
